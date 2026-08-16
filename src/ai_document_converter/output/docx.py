"""خروجی Word ساختاریافته و قابل ویرایش."""

from __future__ import annotations

from pathlib import Path

from ai_document_converter.document.model import BlockType, DocumentModel


class DOCXWriter:
    """مدل ساختاری سند را به DOCX تبدیل می‌کند."""

    name = "docx"

    def write(self, text: str, target: Path, rtl: bool = True) -> Path:
        """سازگاری با مسیر قدیمی متن خام."""
        from ai_document_converter.document.analyzer import DocumentStructureAnalyzer

        model = DocumentStructureAnalyzer().analyze(text, "fas" if rtl else "eng")
        return self.write_model(model, target)

    def write_model(self, model: DocumentModel, target: Path) -> Path:
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError as exc:
            raise RuntimeError("کتابخانه python-docx نصب نشده است.") from exc

        document = Document()
        for block in model.blocks:
            if not block.text.strip():
                continue
            if block.type == BlockType.TITLE:
                paragraph = document.add_heading(block.text, level=min(block.level + 1, 9))
            elif block.type == BlockType.LIST:
                paragraph = document.add_paragraph(block.text, style="List Bullet")
            else:
                paragraph = document.add_paragraph(block.text)

            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.RIGHT if model.direction == "rtl" else WD_ALIGN_PARAGRAPH.LEFT
            )

        document.save(str(target))
        return target
